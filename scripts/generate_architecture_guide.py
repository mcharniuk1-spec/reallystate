from __future__ import annotations

from pathlib import Path
import textwrap
import zipfile
from datetime import datetime, timezone
from xml.sax.saxutils import escape


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "project-architecture-execution-guide.md"
EXPORT_DIR = ROOT / "docs" / "exports"
EXPORT_MD = EXPORT_DIR / "project-architecture-execution-guide.md"
EXPORT_DOCX = EXPORT_DIR / "project-architecture-execution-guide.docx"
EXPORT_PDF = EXPORT_DIR / "project-architecture-execution-guide.pdf"


def _escape_pdf_text(text: str) -> str:
    return text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")


def _markdown_to_lines(markdown: str, width: int = 95) -> list[str]:
    lines: list[str] = []
    for raw in markdown.splitlines():
        line = raw.rstrip()
        if not line:
            lines.append("")
            continue
        if line.startswith("```"):
            continue
        if line.startswith("# "):
            lines.append(line[2:].upper())
            lines.append("")
            continue
        if line.startswith("## "):
            lines.append(line[3:])
            lines.append("-" * min(len(line[3:]), width))
            continue
        if line.startswith("### "):
            lines.append(line[4:])
            continue
        bullet_prefix = ""
        content = line
        if line.startswith("- "):
            bullet_prefix = "- "
            content = line[2:]
        elif line[:2].isdigit() and line[2:4] == ". ":
            bullet_prefix = line[:4]
            content = line[4:]
        wrapped = textwrap.wrap(content, width=max(20, width - len(bullet_prefix))) or [""]
        for idx, part in enumerate(wrapped):
            prefix = bullet_prefix if idx == 0 else " " * len(bullet_prefix)
            lines.append(prefix + part)
    return lines


def _markdown_blocks(markdown: str) -> list[tuple[str, str | list[list[str]]]]:
    """Parse markdown into typed blocks including tables."""
    blocks: list[tuple[str, str | list[list[str]]]] = []
    table_rows: list[list[str]] = []
    in_table = False

    for raw in markdown.splitlines():
        line = raw.rstrip()

        is_table_line = line.startswith("|") and line.endswith("|")
        is_separator = is_table_line and all(c in "|-: " for c in line.replace("|", ""))

        if is_table_line:
            if is_separator:
                in_table = True
                continue
            cells = [c.strip() for c in line.split("|")[1:-1]]
            if not in_table:
                table_rows = [cells]
                in_table = True
            else:
                table_rows.append(cells)
            continue
        else:
            if in_table and table_rows:
                blocks.append(("table", table_rows))
                table_rows = []
            in_table = False

        if not line:
            continue
        if line.startswith("```"):
            continue
        if line.startswith("# "):
            blocks.append(("title", line[2:]))
        elif line.startswith("## "):
            blocks.append(("h1", line[3:]))
        elif line.startswith("### "):
            blocks.append(("h2", line[4:]))
        elif line.startswith("#### "):
            blocks.append(("h3", line[5:]))
        elif line.startswith("- "):
            blocks.append(("bullet", line[2:]))
        elif line[:2].isdigit() and len(line) > 3 and line[2:4] == ". ":
            blocks.append(("number", line[4:]))
        elif line[:1].isdigit() and len(line) > 2 and line[1:3] == ". ":
            blocks.append(("number", line[3:]))
        else:
            blocks.append(("p", line))

    if in_table and table_rows:
        blocks.append(("table", table_rows))

    return blocks


def _write_docx_fallback(path: Path, markdown: str) -> None:
    def paragraph(text: str, style: str | None = None) -> str:
        style_xml = f'<w:pPr><w:pStyle w:val="{style}"/></w:pPr>' if style else ""
        return f'<w:p>{style_xml}<w:r><w:t xml:space="preserve">{escape(text)}</w:t></w:r></w:p>'

    body: list[str] = []
    for kind, text in _markdown_blocks(markdown):
        if kind == "table":
            assert isinstance(text, list)
            rows = text
            if not rows:
                continue
            tbl = '<w:tbl><w:tblPr><w:tblBorders><w:top w:val="single" w:sz="4" w:space="0" w:color="999999"/><w:left w:val="single" w:sz="4" w:space="0" w:color="999999"/><w:bottom w:val="single" w:sz="4" w:space="0" w:color="999999"/><w:right w:val="single" w:sz="4" w:space="0" w:color="999999"/><w:insideH w:val="single" w:sz="4" w:space="0" w:color="999999"/><w:insideV w:val="single" w:sz="4" w:space="0" w:color="999999"/></w:tblBorders><w:tblW w:w="0" w:type="auto"/></w:tblPr>'
            for ri, row in enumerate(rows):
                tbl += '<w:tr>'
                for cell in row:
                    bold = '<w:b/>' if ri == 0 else ''
                    shading = ' <w:shd w:val="clear" w:fill="F3F4F6"/>' if ri == 0 else ''
                    tbl += f'<w:tc><w:tcPr>{shading}</w:tcPr><w:p><w:pPr><w:rPr>{bold}<w:sz w:val="18"/></w:rPr></w:pPr><w:r><w:rPr>{bold}<w:sz w:val="18"/></w:rPr><w:t xml:space="preserve">{escape(cell)}</w:t></w:r></w:p></w:tc>'
                tbl += '</w:tr>'
            tbl += '</w:tbl>'
            body.append(tbl)
        elif kind == "title":
            assert isinstance(text, str)
            body.append(paragraph(text, "Title"))
        elif kind == "h1":
            assert isinstance(text, str)
            body.append(paragraph(text, "Heading1"))
        elif kind == "h2":
            assert isinstance(text, str)
            body.append(paragraph(text, "Heading2"))
        elif kind == "h3":
            assert isinstance(text, str)
            body.append(paragraph(text, "Heading2"))
        elif kind == "bullet":
            assert isinstance(text, str)
            body.append(paragraph("- " + text, "ListParagraph"))
        elif kind == "number":
            assert isinstance(text, str)
            body.append(paragraph("1) " + text, "ListParagraph"))
        else:
            assert isinstance(text, str)
            body.append(paragraph(text))
    document_xml = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:body>{''.join(body)}<w:sectPr><w:pgSz w:w="12240" w:h="15840"/><w:pgMar w:top="720" w:right="720" w:bottom="720" w:left="720" w:header="360" w:footer="360" w:gutter="0"/></w:sectPr></w:body>
</w:document>"""
    styles_xml = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:styles xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:style w:type="paragraph" w:styleId="Normal"><w:name w:val="Normal"/><w:rPr><w:sz w:val="21"/></w:rPr></w:style>
  <w:style w:type="paragraph" w:styleId="Title"><w:name w:val="Title"/><w:rPr><w:b/><w:sz w:val="34"/></w:rPr></w:style>
  <w:style w:type="paragraph" w:styleId="Heading1"><w:name w:val="Heading 1"/><w:rPr><w:b/><w:sz w:val="28"/></w:rPr></w:style>
  <w:style w:type="paragraph" w:styleId="Heading2"><w:name w:val="Heading 2"/><w:rPr><w:b/><w:sz w:val="24"/></w:rPr></w:style>
  <w:style w:type="paragraph" w:styleId="ListParagraph"><w:name w:val="List Paragraph"/><w:pPr><w:ind w:left="360"/></w:pPr><w:rPr><w:sz w:val="20"/></w:rPr></w:style>
</w:styles>"""
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr(
            "[Content_Types].xml",
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?><Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types"><Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/><Default Extension="xml" ContentType="application/xml"/><Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/><Override PartName="/word/styles.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.styles+xml"/><Override PartName="/docProps/core.xml" ContentType="application/vnd.openxmlformats-package.core-properties+xml"/><Override PartName="/docProps/app.xml" ContentType="application/vnd.openxmlformats-officedocument.extended-properties+xml"/></Types>',
        )
        zf.writestr(
            "_rels/.rels",
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?><Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"><Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/><Relationship Id="rId2" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/styles" Target="word/styles.xml"/></Relationships>',
        )
        zf.writestr(
            "word/_rels/document.xml.rels",
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?><Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"><Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/styles" Target="styles.xml"/></Relationships>',
        )
        zf.writestr("word/document.xml", document_xml)
        zf.writestr("word/styles.xml", styles_xml)
        zf.writestr(
            "docProps/core.xml",
            f'<?xml version="1.0" encoding="UTF-8" standalone="yes"?><cp:coreProperties xmlns:cp="http://schemas.openxmlformats.org/package/2006/metadata/core-properties" xmlns:dc="http://purl.org/dc/elements/1.1/" xmlns:dcterms="http://purl.org/dc/terms/" xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance"><dc:title>Project Architecture Execution Guide</dc:title><dc:creator>Codex</dc:creator><dcterms:created xsi:type="dcterms:W3CDTF">{datetime.now(timezone.utc).isoformat()}</dcterms:created></cp:coreProperties>',
        )
        zf.writestr(
            "docProps/app.xml",
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?><Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/extended-properties"><Application>Codex</Application></Properties>',
        )


def _write_docx(path: Path, markdown: str) -> str:
    try:
        from docx import Document  # type: ignore
        from docx.shared import Pt  # type: ignore
        from docx.enum.table import WD_TABLE_ALIGNMENT  # type: ignore
    except Exception:
        _write_docx_fallback(path, markdown)
        return "fallback"

    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(10)

    doc.add_heading("Bulgaria Real Estate Platform Architecture and Execution Guide", level=0)

    for kind, text in _markdown_blocks(markdown):
        if kind == "table":
            assert isinstance(text, list)
            rows = text
            if not rows:
                continue
            num_cols = max(len(r) for r in rows)
            table = doc.add_table(rows=0, cols=num_cols)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
            for ri, row in enumerate(rows):
                cells = table.add_row().cells
                for ci, cell_text in enumerate(row):
                    if ci < len(cells):
                        cells[ci].text = cell_text
                        for paragraph in cells[ci].paragraphs:
                            paragraph.style.font.size = Pt(9)
                            if ri == 0:
                                for run in paragraph.runs:
                                    run.bold = True
            doc.add_paragraph("")
        elif kind == "title":
            continue
        elif kind == "h1":
            assert isinstance(text, str)
            doc.add_heading(text, level=1)
        elif kind == "h2":
            assert isinstance(text, str)
            doc.add_heading(text, level=2)
        elif kind == "h3":
            assert isinstance(text, str)
            doc.add_heading(text, level=3)
        elif kind == "bullet":
            assert isinstance(text, str)
            doc.add_paragraph(text, style="List Bullet")
        elif kind == "number":
            assert isinstance(text, str)
            doc.add_paragraph(text, style="List Number")
        else:
            assert isinstance(text, str)
            doc.add_paragraph(text)
    doc.save(path)
    return "python-docx"


def _build_pdf(lines: list[str]) -> bytes:
    page_width = 612
    page_height = 792
    margin_left = 54
    margin_top = 54
    font_size = 10
    leading = 13
    max_lines = int((page_height - 2 * margin_top) / leading)

    pages = [lines[i : i + max_lines] for i in range(0, len(lines), max_lines)] or [[]]
    objects: list[bytes] = []

    def add_object(payload: str | bytes) -> int:
        if isinstance(payload, str):
            payload = payload.encode("utf-8")
        objects.append(payload)
        return len(objects)

    font_obj = add_object("<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")

    page_obj_ids: list[int] = []
    content_obj_ids: list[int] = []
    pages_obj_placeholder = add_object("<< >>")

    for page_number, page_lines in enumerate(pages, start=1):
        content_lines = ["BT", f"/F1 {font_size} Tf", f"{margin_left} {page_height - margin_top} Td", f"{leading} TL"]
        for line in page_lines:
            content_lines.append(f"({_escape_pdf_text(line)}) Tj")
            content_lines.append("T*")
        content_lines.append(f"(Page {page_number} of {len(pages)}) Tj")
        content_lines.append("ET")
        stream = "\n".join(content_lines).encode("utf-8")
        content_obj_ids.append(
            add_object(b"<< /Length " + str(len(stream)).encode("ascii") + b" >>\nstream\n" + stream + b"\nendstream")
        )
        page_obj_ids.append(
            add_object(
                f"<< /Type /Page /Parent {pages_obj_placeholder} 0 R /MediaBox [0 0 {page_width} {page_height}] "
                f"/Resources << /Font << /F1 {font_obj} 0 R >> >> /Contents {content_obj_ids[-1]} 0 R >>"
            )
        )

    kids = " ".join(f"{page_id} 0 R" for page_id in page_obj_ids)
    objects[pages_obj_placeholder - 1] = (
        f"<< /Type /Pages /Count {len(page_obj_ids)} /Kids [ {kids} ] >>".encode("utf-8")
    )
    catalog_obj = add_object(f"<< /Type /Catalog /Pages {pages_obj_placeholder} 0 R >>")

    pdf = bytearray(b"%PDF-1.4\n%\xe2\xe3\xcf\xd3\n")
    offsets = [0]
    for index, obj in enumerate(objects, start=1):
        offsets.append(len(pdf))
        pdf.extend(f"{index} 0 obj\n".encode("ascii"))
        pdf.extend(obj)
        pdf.extend(b"\nendobj\n")
    xref_offset = len(pdf)
    pdf.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    pdf.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        pdf.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
    pdf.extend(
        (
            f"trailer\n<< /Size {len(objects) + 1} /Root {catalog_obj} 0 R >>\nstartxref\n{xref_offset}\n%%EOF\n"
        ).encode("ascii")
    )
    return bytes(pdf)


def _write_pdf(path: Path, markdown: str) -> str:
    try:
        from reportlab.lib.pagesizes import LETTER  # type: ignore
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle  # type: ignore
        from reportlab.lib.units import inch  # type: ignore
        from reportlab.lib import colors  # type: ignore
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle  # type: ignore
    except Exception:
        path.write_bytes(_build_pdf(_markdown_to_lines(markdown)))
        return "fallback"

    styles = getSampleStyleSheet()
    cell_style = ParagraphStyle("CellStyle", parent=styles["BodyText"], fontSize=8, leading=10)
    header_style = ParagraphStyle("HeaderStyle", parent=styles["BodyText"], fontSize=8, leading=10, fontName="Helvetica-Bold")
    story: list = []

    for kind, text in _markdown_blocks(markdown):
        if kind == "table":
            assert isinstance(text, list)
            rows = text
            if not rows:
                continue
            table_data = []
            for ri, row in enumerate(rows):
                style_to_use = header_style if ri == 0 else cell_style
                table_data.append([Paragraph(escape(c), style_to_use) for c in row])
            num_cols = max(len(r) for r in table_data)
            col_width = (7.0 * inch) / max(num_cols, 1)
            t = Table(table_data, colWidths=[col_width] * num_cols, repeatRows=1)
            t.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.Color(0.95, 0.96, 0.97)),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.Color(0.8, 0.8, 0.8)),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.Color(0.98, 0.98, 0.99)]),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
            ]))
            story.append(t)
            story.append(Spacer(1, 8))
        elif kind == "title":
            assert isinstance(text, str)
            story.append(Paragraph(f"<b>{escape(text)}</b>", styles["Title"]))
            story.append(Spacer(1, 8))
        elif kind == "h1":
            assert isinstance(text, str)
            story.append(Spacer(1, 6))
            story.append(Paragraph(f"<b>{escape(text)}</b>", styles["Heading1"]))
        elif kind == "h2":
            assert isinstance(text, str)
            story.append(Paragraph(f"<b>{escape(text)}</b>", styles["Heading2"]))
        elif kind == "h3":
            assert isinstance(text, str)
            story.append(Paragraph(f"<b>{escape(text)}</b>", styles["Heading3"]))
        elif kind == "bullet":
            assert isinstance(text, str)
            story.append(Paragraph(f"&bull; {escape(text)}", styles["BodyText"]))
        elif kind == "number":
            assert isinstance(text, str)
            story.append(Paragraph(escape(text), styles["BodyText"]))
        else:
            assert isinstance(text, str)
            story.append(Paragraph(escape(text), styles["BodyText"]))
    doc = SimpleDocTemplate(str(path), pagesize=LETTER, title="Project Architecture Execution Guide",
                            leftMargin=0.6*inch, rightMargin=0.6*inch, topMargin=0.6*inch, bottomMargin=0.6*inch)
    doc.build(story)
    return "reportlab"


def main() -> None:
    EXPORT_DIR.mkdir(parents=True, exist_ok=True)
    markdown = SOURCE.read_text(encoding="utf-8")
    EXPORT_MD.write_text(markdown, encoding="utf-8")
    docx_mode = _write_docx(EXPORT_DOCX, markdown)
    pdf_mode = _write_pdf(EXPORT_PDF, markdown)
    print(f"generated architecture guide markdown: {EXPORT_MD}")
    print(f"generated architecture guide docx ({docx_mode}): {EXPORT_DOCX}")
    print(f"generated architecture guide pdf ({pdf_mode}): {EXPORT_PDF}")


if __name__ == "__main__":
    main()
